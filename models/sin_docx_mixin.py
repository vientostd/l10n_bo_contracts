# -*- coding: utf-8 -*-
"""Mixin para generar documentos desde plantillas .docx (Word) con tokens.

Flujo:
1. Se toma la plantilla .docx cargada en `sin.document.template.template_file`.
2. Se reemplazan los tokens {company.name}, {partner.name}, {product_lines}, etc.
   en párrafos, tablas (incluidas anidadas) y encabezados/pies de página.
3. Se convierte el .docx a PDF con LibreOffice headless (soffice --convert-to pdf).
4. El PDF (base64) se guarda en `document_pdf` del contrato/certificado.

Notas:
- Los tokens partidos entre varios runs (muy común al editar en Word) se resuelven
  fusionando el texto del párrafo en un único run conservando el formato del primero.
- El salto de línea '\n' dentro de un valor (p. ej. {product_lines}) se convierte
  en saltos de línea reales dentro del párrafo (w:br).
"""
import base64
import copy
import logging
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)

_TOKEN_RE = re.compile(r'\{([a-zA-Z0-9_.]+)\}')


def _sanitize_filename(name):
    """Normaliza un nombre para usarlo como nombre base de archivo."""
    name = str(name or 'documento').replace('/', '_').replace('\\', '_')
    name = re.sub(r'[^\w\-\s\.]', '_', name)
    name = re.sub(r'\s+', ' ', name).strip()
    return (name or 'documento')[:80]


class SinDocxMixin:
    """Provee el render de plantillas .docx a PDF. Los modelos que lo usan deben
    heredar de este mixin antes de models.Model."""

    # ------------------------------------------------------------------
    # API pública
    # ------------------------------------------------------------------
    def _render_docx_pdf(self, template_b64, ctx, out_name='documento'):
        """Rellena los tokens del .docx contenido en `template_b64` y lo
        convierte a PDF. Devuelve el PDF en base64 (str).

        IMPORTANTE: LibreOffice genera el PDF con el mismo nombre base del
        archivo fuente, por lo que el .docx debe llamarse igual que el PDF
        esperado (<out_name>.pdf)."""
        if not template_b64:
            raise UserError('No hay plantilla .docx cargada en la plantilla.')
        soffice = self._find_soffice()
        safe_name = _sanitize_filename(out_name)
        tmp_dir = tempfile.mkdtemp(prefix='sin_docx_')
        try:
            src = os.path.join(tmp_dir, '%s.docx' % safe_name)
            with open(src, 'wb') as f:
                f.write(base64.b64decode(template_b64))

            self._render_docx_file(src, ctx)

            pdf_path = os.path.join(tmp_dir, '%s.pdf' % safe_name)
            self._convert_to_pdf(soffice, tmp_dir, src, pdf_path)

            with open(pdf_path, 'rb') as f:
                return base64.b64encode(f.read()).decode('utf-8')
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    # ------------------------------------------------------------------
    # Render de tokens en el documento
    # ------------------------------------------------------------------
    def _render_docx_file(self, docx_path, ctx):
        """Reemplaza tokens en párrafos, tablas y encabezados/pies del docx."""
        import docx
        document = docx.Document(docx_path)

        # Cuerpo principal
        for paragraph in document.paragraphs:
            self._fill_paragraph(paragraph, ctx)
        self._fill_tables_recursive(document, ctx)

        # Encabezados y pies de todas las secciones
        for section in document.sections:
            for hf in (section.header, section.footer):
                if not hf.is_linked_to_previous:
                    for paragraph in hf.paragraphs:
                        self._fill_paragraph(paragraph, ctx)
                    self._fill_tables_recursive(hf, ctx)
            if section.different_first_page_header_footer:
                for hf in (section.first_page_header, section.first_page_footer):
                    if not hf.is_linked_to_previous:
                        for paragraph in hf.paragraphs:
                            self._fill_paragraph(paragraph, ctx)
                        self._fill_tables_recursive(hf, ctx)

        document.save(docx_path)
        return docx_path

    def _fill_tables_recursive(self, tables_owner, ctx):
        """Recorre tablas (incluidas las anidadas dentro de celdas)."""
        for table in tables_owner.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._fill_paragraph(paragraph, ctx)
                    # tablas anidadas dentro de la celda
                    self._fill_tables_recursive(cell, ctx)

    def _fill_paragraph(self, paragraph, ctx):
        """Reemplaza tokens en un párrafo. Devuelve True si hubo cambios."""
        if not paragraph.runs:
            return False
        changed = False
        # 1er intento: token completo dentro de un único run (caso habitual).
        # El setter run.text conserva el formato (rPr) del run y convierte
        # '\n' en saltos de línea reales.
        for run in paragraph.runs:
            text = run.text or ''
            if '{' in text:
                new_text = self._render_plain_tokens(text, ctx)
                if new_text != text:
                    run.text = new_text
                    changed = True
        # 2do intento: token partido entre runs (p. ej. "{company." + "name}").
        # Se fusiona el párrafo en un único run conservando el rPr del primero.
        full = ''.join(r.text or '' for r in paragraph.runs)
        if '{' in full:
            new_full = self._render_plain_tokens(full, ctx)
            if new_full != full:
                self._merge_paragraph_text(paragraph, new_full)
                changed = True
        return changed

    def _merge_paragraph_text(self, paragraph, text):
        """Sustituye todos los runs del párrafo por uno solo con `text`,
        conservando el rPr (formato) del primer run que lo tenga."""
        rpr = None
        for run in paragraph.runs:
            if run.text:
                if run._r.rPr is not None:
                    rpr = copy.deepcopy(run._r.rPr)
                break
        for run in list(paragraph.runs):
            run._r.getparent().remove(run._r)
        run = paragraph.add_run()
        if rpr is not None:
            run._r.insert(0, rpr)
        run.text = text

    @staticmethod
    def _render_plain_tokens(text, ctx):
        """Reemplaza {token} en texto plano (sin HTML). Igual regex que el
        flujo HTML para mantener los mismos tokens disponibles."""
        def _sub(m):
            key = m.group(1)
            return str(ctx.get(key, m.group(0)))
        return _TOKEN_RE.sub(_sub, text)

    # ------------------------------------------------------------------
    # Conversión a PDF con LibreOffice
    # ------------------------------------------------------------------
    def _find_soffice(self):
        for name in ('soffice', 'libreoffice'):
            path = shutil.which(name)
            if path:
                return path
        for candidate in (
            '/usr/bin/soffice',
            '/usr/lib/libreoffice/program/soffice',
            '/opt/libreoffice/program/soffice',
            '/usr/local/bin/soffice',
        ):
            if os.path.exists(candidate):
                return candidate
        raise UserError(
            'LibreOffice no está instalado en el servidor. Ejecute como root:\n'
            '  apt-get install -y --no-install-recommends libreoffice-writer'
        )

    def _convert_to_pdf(self, soffice, out_dir, src_path, pdf_path):
        """Convierte el .docx a PDF con LibreOffice headless."""
        # Perfil temporal aislado por conversión para evitar bloqueos/conflictos
        # entre conversiones concurrentes (odeo con workers paralelos).
        profile_uri = Path(os.path.join(out_dir, 'lo_profile')).as_uri()
        cmd = [
            soffice,
            '--headless',
            '--norestore',
            '--nologo',
            '-env:UserInstallation=%s' % profile_uri,
            '--convert-to', 'pdf',
            '--outdir', out_dir,
            src_path,
        ]
        try:
            result = subprocess.run(
                cmd, capture_output=True, text=True, timeout=180,
            )
        except subprocess.TimeoutExpired:
            raise UserError(
                'La conversión a PDF excedió el tiempo máximo (180 s). '
                'Revise que LibreOffice esté operativo en el servidor.'
            )
        if result.returncode != 0:
            _logger.error('soffice falló: %s', result.stderr)
            raise UserError(
                'LibreOffice no pudo convertir el documento a PDF.\n'
                'Detalle: %s' % (result.stderr or 'código de salida %s' % result.returncode)
            )
        if not os.path.exists(pdf_path):
            raise UserError('LibreOffice no generó el archivo PDF esperado.')
        return pdf_path